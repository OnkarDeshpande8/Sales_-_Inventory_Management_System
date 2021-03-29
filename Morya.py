#   Sales & Inventory Management System for MORYA SALES AND SERVICES 
#   Author = Onkar Anil Deshpande 
#   Date = 24/01/2021
#-------------------------------------------------------------------------------------------------------------------------

import time
from tkinter import *
from tkinter.ttk import * 
from PIL import ImageTk,Image
import pandas as pds
from pandas import ExcelWriter
from pandas import ExcelFile
import numpy
import datetime
import docx
from docx import Document                  
from docx.shared import Inches,Pt
from docx.shared import RGBColor
import os
import gspread 
import pygsheets
from oauth2client.service_account import ServiceAccountCredentials
import logging

Debug_log = "Debug.log"
Root_Geometry = '1366x768'
Root_Title='MORYA SALES AND SERVICES '

BG_WHITE = "white"
BG_DeepSkyBlue= "deep sky blue"
BG_LightSkyBlue="light sky blue"
FG_Black="black"
FG_Red ="red"
FG_Green = 'Green'
Button_Style = 'W.TButton'

Header_Font1=('calibri',25,'bold')
Label_Font1=('Helvetica',20,'bold')
Label_Font2=('calibri',20,'bold')
Label_Font3=('calibre',15,'bold')
Label_Font4=('calibre',10,'bold')
Label_Font5=('calibre',20,'bold')
Label_Font6=('Helvetica',15,'bold')
Label_Font7=('calibri',15,'italic')
Label_Font8=("Helvetica",30,'bold')
Label_Font9=("Helvetica",25,'bold')
Entry_Font1 =('calibre',15,'normal')

ImagePath = "mm2.png"

TYRE_SIZE = "TYRE SIZE"
CEAT_TYRE="CEAT TYRE"
APOLLO_TYRE='APOLLO TYRE'
NUEMEX_TYRE='NUEMEX TYRE'
TVS_TYRE='TVS TYRE'
STELLBIRID_TYRE='STELLBIRID TYRE'
MRF_TYRE='MRF TYRE'
METRO_TYRE='METRO TYRE'

file1=('Excel_files\\CEAT TYRE.xlsx')
file2=('Excel_files\\APOLLO TYRE.xlsx')
file3=('Excel_files\\NUEMEX TYRE.xlsx')
file4=('Excel_files\\TVS TYRE.xlsx')
file5=('Excel_files\\STELLBIRID TYRE.xlsx')
file6=('Excel_files\\MRF TYRE.xlsx')
file7=('Excel_files\\METRO TYRE.xlsx') 
file8=('Excel_files\\Sales_report.xlsx')
file9=('Excel_files\\tube.xlsx') 
file10=('Excel_files\\Sales_report_tube.xlsx')

STATUS_ERROR1 = "Error : "

def tyre():
    A = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
    A.place(x=0,y=195)
    B = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
    B.place(x=0,y=195)
    c = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1166,relief=FLAT) 
    c.place(x=202,y=195)
    header0 = Label(text="TYRE SECTION", font = Header_Font1,background=BG_DeepSkyBlue)
    header0.place(x=690,y=197)
    refresh_tyre_button = Button(root, text = "Refresh",style = Button_Style,command = tyre).place(x=1200,y=600)
   
    def inward_stock1():
        A = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
        A.place(x=0,y=195)
        B = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
        B.place(x=0,y=195)
        c = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1166,relief=FLAT) 
        c.place(x=202,y=195)
        header0 = Label(text="TYRE INWARD SECTION", font = Header_Font1,background=BG_DeepSkyBlue)
        header0.place(x=630,y=197)
        refresh_tyreInward_button = Button(root, text = "Refresh",style = Button_Style,command = inward_stock1).place(x=1200,y=600)

        inward_stock = Button(root, text = "Inward Stock",style = Button_Style,command=inward_stock1)
        inward_stock.place(x=40,y=245)
        
        header1 = Label(root, text="Select Brand", font = Label_Font2,background=BG_WHITE)
        header2 = Label(root, text="Select Tyre Size", font = Label_Font2,background=BG_WHITE)
        tt_tl = Label(text="TT / TL  = ",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        stock = Label(text=" Stock   =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        ndp = Label(text=  "  NDP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        mrp = Label(text=  "  MRP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        quntity = Label(text="Quantity",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)

        qunt_var =StringVar() 
        quntity_entry = Entry(root, textvariable = qunt_var,font=Entry_Font1,width=10)
        add_button = Button(root, text = "Add Inward Stock",style = Button_Style)
        newstock=0

        header1.place(x=300,y=250)
        header2.place(x=550,y=250)
        tt_tl.place(x = 300, y =350)
        stock.place(x = 300, y =430)
        ndp.place(x = 300, y =510)
        mrp.place(x = 300, y =590)
        quntity.place(x = 770, y =350)
        quntity_entry.place(x = 950, y =355)

        mainframe3 = Frame(root)
        mainframe3.place(x=300,y=290)
        tkvar3 = StringVar(root)
        popupMenu3 = OptionMenu(mainframe3, tkvar3, *Tyre_Brand_Choice)
        popupMenu3.configure(width=22)
        popupMenu3.pack()

        def add(var1):
            add_fail1 = Label(text = 'Enter Quntity', font=Label_Font4,background=BG_WHITE, foreground = FG_Red)
            add_sucess = Label(text = 'Successful', font=Label_Font5,background=BG_WHITE, foreground = FG_Green)
            if quntity_entry.get() == "":
                add_fail1.place(x=885,y=395)
            else:
                add_quntity = int(quntity_entry.get())
                newstock=var1+add_quntity
                print(add_quntity)
                print(newstock)
                add_fail1.destroy()
                add_sucess.place(x=865,y=390)
            return newstock

        def change_dropdown(*args):
            mainframe4 = Frame(root)
            mainframe4.place(x=554,y=290)
            tkvar1 = StringVar(root)

            my_list1= StringVar(root)

            tt_tl1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            stock1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            ndp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            mrp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            add_fail1 = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            add_sucess = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
            tt_tl1.place(x = 450, y =350)
            stock1.place(x = 450, y =430)
            ndp1.place(x = 450, y =510)
            mrp1.place(x = 450, y =590)
            add_fail1.place(x=885,y=395)
            add_sucess.place(x=865,y=390)
        
            if tkvar3.get() == CEAT_TYRE:
                sheet1 = pds.read_excel(file1, sheet_name = CEAT_TYRE) 
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices8 = sheet1[TYRE_SIZE].tolist()
                popupMenu8 = OptionMenu(mainframe4, tkvar1, *choices8)
                popupMenu8.configure(width=22)
                popupMenu8.pack()

                def change_dropdown1(*args):
                    sheet1.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet1.loc[tkvar1.get()][1]
                    my_list2=sheet1.loc[tkvar1.get()][2]
                    my_list3=sheet1.loc[tkvar1.get()][3]
                    my_list4=sheet1.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)
                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet1.loc[tkvar1.get(),'Stock']=newstock
                        sheet1.to_excel(file1, sheet_name = CEAT_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
                
                tkvar1.trace('w', change_dropdown1) 
        
            if tkvar3.get() == APOLLO_TYRE:
                sheet2 = pds.read_excel(file2, sheet_name = APOLLO_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices9 = sheet2[TYRE_SIZE].tolist()
                popupMenu9 = OptionMenu(mainframe4, tkvar1, *choices9)
                popupMenu9.configure(width=22)
                popupMenu9.pack()

                def change_dropdown2(*args):
                    sheet2.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet2.loc[tkvar1.get()][1]
                    my_list2=sheet2.loc[tkvar1.get()][2]
                    my_list3=sheet2.loc[tkvar1.get()][3]
                    my_list4=sheet2.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)
                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet2.loc[tkvar1.get(),'Stock']=newstock
                        sheet2.to_excel(file2, sheet_name = APOLLO_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
                         
                tkvar1.trace('w', change_dropdown2)

            if tkvar3.get() == NUEMEX_TYRE:
                sheet3 = pds.read_excel(file3, sheet_name = NUEMEX_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices10 = sheet3[TYRE_SIZE].tolist()
                popupMenu10 = OptionMenu(mainframe4, tkvar1, *choices10)
                popupMenu10.configure(width=22)
                popupMenu10.pack()

                def change_dropdown3(*args):
                    sheet3.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet3.loc[tkvar1.get()][1]
                    my_list2=sheet3.loc[tkvar1.get()][2]
                    my_list3=sheet3.loc[tkvar1.get()][3]
                    my_list4=sheet3.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)
                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet3.loc[tkvar1.get(),'Stock']=newstock
                        sheet3.to_excel(file3, sheet_name = NUEMEX_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
 
                tkvar1.trace('w', change_dropdown3)

            if tkvar3.get() == TVS_TYRE:
                sheet4 = pds.read_excel(file4, sheet_name = TVS_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices11 = sheet4[TYRE_SIZE].tolist()
                popupMenu11 = OptionMenu(mainframe4, tkvar1, *choices11)
                popupMenu11.configure(width=22)
                popupMenu11.pack()

                def change_dropdown4(*args):
                    sheet4.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet4.loc[tkvar1.get()][1]
                    my_list2=sheet4.loc[tkvar1.get()][2]
                    my_list3=sheet4.loc[tkvar1.get()][3]
                    my_list4=sheet4.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)

                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet4.loc[tkvar1.get(),'Stock']=newstock
                        sheet4.to_excel(file4, sheet_name = TVS_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
                        
                tkvar1.trace('w', change_dropdown4)

            if tkvar3.get() == STELLBIRID_TYRE:
                sheet5 = pds.read_excel(file5, sheet_name = STELLBIRID_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices12 = sheet5[TYRE_SIZE].tolist()
                popupMenu12 = OptionMenu(mainframe4, tkvar1, *choices12)
                popupMenu12.configure(width=22)
                popupMenu12.pack()

                def change_dropdown5(*args):
                    sheet5.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet5.loc[tkvar1.get()][1]
                    my_list2=sheet5.loc[tkvar1.get()][2]
                    my_list3=sheet5.loc[tkvar1.get()][3]
                    my_list4=sheet5.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)

                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet5.loc[tkvar1.get(),'Stock']=newstock
                        sheet5.to_excel(file5, sheet_name = STELLBIRID_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
                
                tkvar1.trace('w', change_dropdown5)
                
            if tkvar3.get() == MRF_TYRE:
                sheet6 = pds.read_excel(file6, sheet_name = MRF_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices13 = sheet6[TYRE_SIZE].tolist()
                popupMenu13 = OptionMenu(mainframe4, tkvar1, *choices13)
                popupMenu13.configure(width=22)
                popupMenu13.pack()

                def change_dropdown6(*args):
                    sheet6.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet6.loc[tkvar1.get()][1]
                    my_list2=sheet6.loc[tkvar1.get()][2]
                    my_list3=sheet6.loc[tkvar1.get()][3]
                    my_list4=sheet6.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)

                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet6.loc[tkvar1.get(),'Stock']=newstock
                        sheet6.to_excel(file6, sheet_name = MRF_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
                
                tkvar1.trace('w', change_dropdown6)

            if tkvar3.get() == METRO_TYRE:
                sheet7 = pds.read_excel(file7, sheet_name = METRO_TYRE)
                mainframe4 = Frame(root)
                mainframe4.place(x=554,y=290)
                tkvar1 = StringVar(root)
                choices14 = sheet7[TYRE_SIZE].tolist()
                popupMenu14 = OptionMenu(mainframe4, tkvar1, *choices14)
                popupMenu14.configure(width=22)
                popupMenu14.pack()

                def change_dropdown7(*args):
                    sheet7.set_index(TYRE_SIZE, inplace = True)
                    my_list1=sheet7.loc[tkvar1.get()][1]
                    my_list2=sheet7.loc[tkvar1.get()][2]
                    my_list3=sheet7.loc[tkvar1.get()][3]
                    my_list4=sheet7.loc[tkvar1.get()][4]
                    print (my_list1 , my_list2 , my_list3 ,my_list4)

                    tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                    stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                    ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                    mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                    def add1():
                        newstock=add(my_list4)
                        print('newstock=',newstock)
                        sheet7.loc[tkvar1.get(),'Stock']=newstock
                        sheet7.to_excel(file7, sheet_name = METRO_TYRE)
                    add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
                    add_button.place(x = 870, y =430)
    
                tkvar1.trace('w', change_dropdown7)

        # link function to change dropdown
        tkvar3.trace('w', change_dropdown)

    inward_stock = Button(root, text = "Inward Stock",style = Button_Style,command=inward_stock1)
    inward_stock.place(x=40,y=245)

    header1 = Label(root, text="Select Brand", font = Label_Font2,background=BG_WHITE)
    header2 = Label(root, text="Select Tyre Size", font = Label_Font2,background=BG_WHITE)
    tt_tl = Label(text="TT / TL  = ",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    stock = Label(text=" Stock   =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    ndp = Label(text=  "  NDP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    mrp = Label(text=  "  MRP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    header1.place(x=300,y=250)
    header2.place(x=550,y=250)
    tt_tl.place(x = 300, y =350)
    stock.place(x = 300, y =430)
    ndp.place(x = 300, y =510)
    mrp.place(x = 300, y =590)

    cust_name = Label(text="Customer Name",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    cust_number = Label(text="Mobile no.",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    quntity = Label(text="Quantity",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    sellmrp = Label(text="Selling Price",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    cust_name.place(x = 780, y =350)
    cust_number.place(x = 780, y =400)
    quntity.place(x = 780, y =450)
    sellmrp.place(x = 780, y =500)

    sell_button = Button(root, text = "Sell",style = Button_Style)
    newstock=0
    profit=0

    mainframe = Frame(root)
    mainframe.place(x=300,y=290)
    tkvar = StringVar(root)
    Tyre_Brand_Choice = ['Select Brand',CEAT_TYRE,APOLLO_TYRE,NUEMEX_TYRE,TVS_TYRE,STELLBIRID_TYRE,MRF_TYRE,METRO_TYRE]
    popupMenu = OptionMenu(mainframe, tkvar, *Tyre_Brand_Choice)
    popupMenu.configure(width=22)
    popupMenu.pack()
    
    # Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
    # df = pds.DataFrame(Sales_report_tyre)
    # no=len(df['No.'])+1
    x = datetime.datetime.now()
    date1=x.strftime("%d/%m/%y")
    date2=x.strftime("%d/%m/%Y")

    def change_dropdown(*args):
        # mainframe1 = Frame(root)
        # mainframe1.place(x=554,y=290)
        # tkvar1 = StringVar(root)
        
        name_var=StringVar()
        mobile_var=StringVar()
        qunt_var =StringVar() 
        sellmrp_var = StringVar()
        cust_name_entry= Entry(root,textvariable=name_var,font=Entry_Font1,width=20)
        cust_number_entry= Entry(root,textvariable=mobile_var,font=Entry_Font1,width=10)
        quntity_entry = Entry(root, textvariable = qunt_var,font=Entry_Font1,width=10)
        sellmrp_entry = Entry(root, textvariable = sellmrp_var,font=Entry_Font1,width=10)
        cust_name_entry.place(x = 950, y =350)
        cust_number_entry.place(x = 950, y =400)
        quntity_entry.place(x = 950, y =450)
        sellmrp_entry.place(x = 950, y =500)


        def print1(var1,var2,var3):
            document = Document()
            
            ab=document.add_paragraph()
            ab.add_run('\t\t\t\t\tINVOICE').bold = True
            ab=document.add_paragraph()
            ab.add_run('Morya Sales & Services').bold = True
            
            a = document.add_paragraph('Main Road, Mane Nagar, Rendal-Hupari\n')
            a.add_run('Tal-Hatkanangle, Dist.-Kolhapur\n')
            a.add_run('Pin-416203\n')
            a.add_run('Phone: 8208600074, 9067173171\n')
            a.add_run('_________________________________________________________________________________________________________').bold = True

            name=str(cust_name_entry.get())
            mono=str(cust_number_entry.get())
            b = document.add_paragraph('\t\t\t\t\t\t\t\t\tDate : ')
            b.add_run(date2)
            b = document.add_paragraph('Customer Name : ')
            b.add_run(name)
            b = document.add_paragraph('Mobile No. : ')
            b.add_run(mono)

            print_total=var2*var3

            recordset = [
                {
                    "No." : 1,
                    "desc": tkvar.get()+' : '+var1,
                    "qty": var2,
                    "Unit_Price": var3,
                    "Total":print_total
                }
            ]
            table = document.add_table(rows=1, cols=5, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'No.'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Quantity'
            hdr_cells[3].text = 'Unit Price'
            hdr_cells[4].text = 'Total'
            for item in recordset:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['No.'])
                row_cells[1].text = str(item['desc'])
                row_cells[2].text = str(item['qty'])
                row_cells[3].text = str(item['Unit_Price'])
                row_cells[4].text = str(item['Total'])

            ac=document.add_paragraph()
            ac.add_run('\n\n\n\n\n\n\n\n\n\n\n\n\t\t\t\t\t\t\t\t   Morya Sales & Services').bold = True
            
            document.save('simple.docx')
            os.startfile('simple.docx','print')
            print('print')
    
        def sell(var1,var2):
            sell_fail1 = Label(text = 'Enter Quntity', font=Label_Font4,background=BG_WHITE, foreground =FG_Red)
            sell_fail = Label(text = 'Enter Selling Price', font=Label_Font4,background=BG_WHITE, foreground =FG_Red)
            sell_sucess = Label(text = 'Successful', font=Label_Font5,background=BG_WHITE, foreground =FG_Green)
            if quntity_entry.get() == "":
                sell_fail1.place(x=885,y=550)
            elif sellmrp_entry.get() == "":
                sell_fail.place(x=880,y=550)
            else:
                sell_quntity = int(quntity_entry.get())
                sell_price = int(sellmrp_entry.get())
                name1=str(cust_name_entry.get())
                number1=int(cust_number_entry.get())
                newstock=var1-sell_quntity
                profit = (sell_price-var2)*sell_quntity
                print(sell_quntity)
                print(sell_price)
                print(newstock)
                print(profit)
                sell_fail.destroy()
                sell_fail1.destroy()
                sell_sucess.place(x=870,y=550)
            return newstock,profit,sell_quntity,sell_price

        my_list1= StringVar(root)

        tt_tl1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        stock1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        ndp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        mrp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_fail = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_fail1 = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_sucess = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)

        tt_tl1.place(x = 450, y =350)
        stock1.place(x = 450, y =430)
        ndp1.place(x = 450, y =510)
        mrp1.place(x = 450, y =590)

        sell_fail1.place(x=880,y=550)
        sell_fail.place(x=880,y=550)
        sell_sucess.place(x=870,y=550)
        
        if tkvar.get() == CEAT_TYRE:
            sheet1 = pds.read_excel(file1, sheet_name = CEAT_TYRE) 
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices1 = sheet1[TYRE_SIZE].tolist()
            popupMenu1 = OptionMenu(mainframe1, tkvar1, *choices1)
            popupMenu1.configure(width=22)
            popupMenu1.pack()

            def change_dropdown1(*args):
                sheet1.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet1.loc[tkvar1.get()][1]
                my_list2=sheet1.loc[tkvar1.get()][2]
                my_list3=sheet1.loc[tkvar1.get()][3]
                my_list4=sheet1.loc[tkvar1.get()][4]
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet1.loc[tkvar1.get(),'Stock']=newstock
                    sheet1.to_excel(file1, sheet_name = CEAT_TYRE)
                    
                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)
                    no=len(df['No.'])+1
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False) 
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)

                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)
            tkvar1.trace('w', change_dropdown1) 
      
        if tkvar.get() == APOLLO_TYRE:
            sheet2 = pds.read_excel(file2, sheet_name = APOLLO_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices2 = sheet2[TYRE_SIZE].tolist()
            popupMenu2 = OptionMenu(mainframe1, tkvar1, *choices2)
            popupMenu2.configure(width=22)
            popupMenu2.pack()

            def change_dropdown2(*args):
                sheet2.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet2.loc[tkvar1.get()][1]
                my_list2=sheet2.loc[tkvar1.get()][2]
                my_list3=sheet2.loc[tkvar1.get()][3]
                my_list4=sheet2.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet2.loc[tkvar1.get(),'Stock']=newstock
                    sheet2.to_excel(file2, sheet_name = APOLLO_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)
                    no=len(df['No.'])+1
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)
            tkvar1.trace('w', change_dropdown2)
               
        if tkvar.get() == NUEMEX_TYRE:
            sheet3 = pds.read_excel(file3, sheet_name = NUEMEX_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices3 = sheet3[TYRE_SIZE].tolist()
            popupMenu3 = OptionMenu(mainframe1, tkvar1, *choices3)
            popupMenu3.configure(width=22)
            popupMenu3.pack()

            def change_dropdown3(*args):
                sheet3.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet3.loc[tkvar1.get()][1]
                my_list2=sheet3.loc[tkvar1.get()][2]
                my_list3=sheet3.loc[tkvar1.get()][3]
                my_list4=sheet3.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet3.loc[tkvar1.get(),'Stock']=newstock
                    sheet3.to_excel(file3, sheet_name = NUEMEX_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)                    
                    no=len(df['No.'])+1
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)

            tkvar1.trace('w', change_dropdown3)

        if tkvar.get() == TVS_TYRE:
            sheet4 = pds.read_excel(file4, sheet_name = TVS_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices4 = sheet4[TYRE_SIZE].tolist()
            popupMenu4 = OptionMenu(mainframe1, tkvar1, *choices4)
            popupMenu4.configure(width=22)
            popupMenu4.pack()

            def change_dropdown4(*args):
                sheet4.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet4.loc[tkvar1.get()][1]
                my_list2=sheet4.loc[tkvar1.get()][2]
                my_list3=sheet4.loc[tkvar1.get()][3]
                my_list4=sheet4.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet4.loc[tkvar1.get(),'Stock']=newstock
                    sheet4.to_excel(file4, sheet_name = TVS_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)   
                    no=len(df['No.'])+1                 
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)
   
            tkvar1.trace('w', change_dropdown4)

        if tkvar.get() == STELLBIRID_TYRE:
            sheet5 = pds.read_excel(file5, sheet_name = STELLBIRID_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices5 = sheet5[TYRE_SIZE].tolist()
            popupMenu5 = OptionMenu(mainframe1, tkvar1, *choices5)
            popupMenu5.configure(width=22)
            popupMenu5.pack()

            def change_dropdown5(*args):
                sheet5.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet5.loc[tkvar1.get()][1]
                my_list2=sheet5.loc[tkvar1.get()][2]
                my_list3=sheet5.loc[tkvar1.get()][3]
                my_list4=sheet5.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet5.loc[tkvar1.get(),'Stock']=newstock
                    sheet5.to_excel(file5, sheet_name = STELLBIRID_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)
                    no=len(df['No.'])+1                    
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)

            tkvar1.trace('w', change_dropdown5)

        if tkvar.get() == MRF_TYRE:
            sheet6 = pds.read_excel(file6, sheet_name = MRF_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices6 = sheet6[TYRE_SIZE].tolist()
            popupMenu6 = OptionMenu(mainframe1, tkvar1, *choices6)
            popupMenu6.configure(width=22)
            popupMenu6.pack()

            def change_dropdown6(*args):
                sheet6.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet6.loc[tkvar1.get()][1]
                my_list2=sheet6.loc[tkvar1.get()][2]
                my_list3=sheet6.loc[tkvar1.get()][3]
                my_list4=sheet6.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet6.loc[tkvar1.get(),'Stock']=newstock
                    sheet6.to_excel(file6, sheet_name = MRF_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)   
                    no=len(df['No.'])+1                 
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)
            
            tkvar1.trace('w', change_dropdown6)

        if tkvar.get() == METRO_TYRE:
            sheet7 = pds.read_excel(file7, sheet_name = METRO_TYRE)
            mainframe1 = Frame(root)
            mainframe1.place(x=554,y=290)
            tkvar1 = StringVar(root)
            choices7 = sheet7[TYRE_SIZE].tolist()
            popupMenu7 = OptionMenu(mainframe1, tkvar1, *choices7)
            popupMenu7.configure(width=22)
            popupMenu7.pack()

            def change_dropdown7(*args):
                sheet7.set_index(TYRE_SIZE, inplace = True)
                my_list1=sheet7.loc[tkvar1.get()][1]
                my_list2=sheet7.loc[tkvar1.get()][2]
                my_list3=sheet7.loc[tkvar1.get()][3]
                my_list4=sheet7.loc[tkvar1.get()][4]
                print (my_list1 , my_list2 , my_list3)
                print(my_list4)
                tt_tl1 = Label(text=my_list1,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
                stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
                ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
                mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =590)
                def sell1():
                    newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
                    name1=str(cust_name_entry.get())
                    number1=int(cust_number_entry.get())
                    print('newstock=',newstock)
                    sheet7.loc[tkvar1.get(),'Stock']=newstock
                    sheet7.to_excel(file7, sheet_name = METRO_TYRE)

                    Sales_report_tyre = pds.read_excel(file8, sheet_name = "Sheet1")
                    df = pds.DataFrame(Sales_report_tyre)  
                    no=len(df['No.'])+1                  
                    df2 = pds.DataFrame({"No.":[no],"Date":[date1],"Tyre":[tkvar.get()],TYRE_SIZE:[tkvar1.get()],"TT/TL":[my_list1], "Quantity":[sell_quntity],"NDP":[my_list2], "MRP":[sell_price], "Profit":[profit],"Customer Name":[name1],"Mobile No.":[number1]})
                    append_sell=df.append(df2,ignore_index = True)
                    append_sell.to_excel(file8, sheet_name = "Sheet1", index=False)  
                    def print2():
                        print1(tkvar1.get(),sell_quntity,sell_price)
                    print_button = Button(root, text = "Print",style = Button_Style,command = print2)
                    print_button.place(x = 880, y =640)
                sell_button = Button(root, text = "Sell",style = Button_Style,command = sell1)
                sell_button.place(x = 880, y =600)
            
            tkvar1.trace('w', change_dropdown7)

    tkvar.trace('w', change_dropdown)

def tubes():
    c = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
    c.place(x=0,y=195)
    BG_LightSkyBlue=BG_DeepSkyBlue
    d = Canvas(root, bg =BG_LightSkyBlue, height = 500, width = 200,relief=FLAT) 
    d.place(x=0,y=195)
    e = Canvas(root, bg =BG_LightSkyBlue, height = 45, width = 1166,relief=FLAT) 
    e.place(x=202,y=195)
    header0 = Label(text="TUBE SECTION", font = Header_Font1,background=BG_LightSkyBlue) 
    header0.place(x=690,y=197)
    refresh_tube_button = Button(root, text = "Refresh",style = Button_Style,command = tubes).place(x=1200,y=600)

    def inward_stock1():
        c = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
        c.place(x=0,y=195)
        d = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
        d.place(x=0,y=195)
        e = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1166,relief=FLAT) 
        e.place(x=202,y=195)
        header0 = Label(text="TUBE INWARD SECTION", font = Header_Font1,background=BG_DeepSkyBlue)
        header0.place(x=630,y=197)
        refresh_ttubeInward_button = Button(root, text = "Refresh",style = Button_Style,command = tubes).place(x=1200,y=600)

        inward_stock = Button(root, text = "Inward Stock",style = Button_Style,command=inward_stock1)
        inward_stock.place(x=40,y=245)
        header3 = Label(root, text="Select Tube Size", font = Label_Font2,background=BG_WHITE)
        stock = Label(text=" Stock   =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        ndp = Label(text=  "  NDP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        mrp = Label(text=  "  MRP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
        quntity = Label(text="Quantity",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)

        qunt_var =StringVar() 
        quntity_entry = Entry(root, textvariable = qunt_var,font=Entry_Font1,width=10)
        add_button = Button(root, text = "Add Inward Stock",style = Button_Style)
        newstock=0
        profit=0

        header3.place(x=300,y=250)
        stock.place(x = 300, y =350)
        ndp.place(x = 300, y =430)
        mrp.place(x = 300, y =510)
        quntity.place(x = 770, y =350)
        quntity_entry.place(x = 950, y =355)

        sheet19 = pds.read_excel(file9, sheet_name = "tube") 
        mainframe9 = Frame(root)
        mainframe9.place(x=300,y=290)
        tkvar9 = StringVar(root)
        choices9 = sheet19[TYRE_SIZE].tolist()
        popupMenu9 = OptionMenu(mainframe9, tkvar9, *choices9)
        popupMenu9.configure(width=22)
        popupMenu9.pack()

        my_list1= StringVar(root)

        stock1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        ndp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        mrp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_fail = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_fail1 = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
        sell_sucess = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)

        stock1.place(x = 450, y =350)
        ndp1.place(x = 450, y =430)
        mrp1.place(x = 450, y =510)
        sell_fail1.place(x=865,y=490)
        sell_fail.place(x=850,y=490)
        sell_sucess.place(x=830,y=480)
        
        def add(var1):
            print('add')
            add_fail1 = Label(text = 'Enter Quntity', font=Label_Font4,background=BG_WHITE, foreground = FG_Red)
            add_sucess = Label(text = 'Successful', font=Label_Font5,background=BG_WHITE, foreground = FG_Green)
            if quntity_entry.get() == "":
                add_fail1.place(x=885,y=395)
            else:
                add_quntity = int(quntity_entry.get())
                newstock=var1+add_quntity
                print(add_quntity)
                print(newstock)
                add_fail1.destroy()
                add_sucess.place(x=865,y=390)
            return newstock
        
        def change_dropdown(*args):
            sheet19.set_index(TYRE_SIZE, inplace = True)
            my_list2=sheet19.loc[tkvar9.get()][1]
            my_list3=sheet19.loc[tkvar9.get()][2]
            my_list4=sheet19.loc[tkvar9.get()][3]
            print (my_list2 , my_list3, my_list4)
            stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
            ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
            mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
            def add1():
                newstock=add(my_list4)
                print('newstock=',newstock)
                sheet19.loc[tkvar9.get(),'Stock']=newstock
                sheet19.to_excel(file9, sheet_name = 'tube')
            add_button = Button(root, text = "Add Inward Stock",style = Button_Style,command = add1)
            add_button.place(x = 870, y =430)      

        tkvar9.trace('w', change_dropdown)

    inward_stock = Button(root, text = "Inward Stock",style = Button_Style,command=inward_stock1)
    inward_stock.place(x=40,y=245)
    header3 = Label(root, text="Select Tube Size", font = Label_Font2,background=BG_WHITE)
    stock = Label(text=" Stock   =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    ndp = Label(text=  "  NDP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    mrp = Label(text=  "  MRP    =",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    header3.place(x=300,y=250)
    stock.place(x = 300, y =350)
    ndp.place(x = 300, y =430)
    mrp.place(x = 300, y =510)
    
    cust_name = Label(text="Customer Name",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    cust_number = Label(text="Mobile no.",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    quntity = Label(text="Quantity",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    sellmrp = Label(text="Selling Price",font=Label_Font6,background=BG_WHITE,foreground=FG_Black)
    cust_name.place(x = 780, y =350)
    cust_number.place(x = 780, y =400)
    quntity.place(x = 780, y =450)
    sellmrp.place(x = 780, y =500)

    name_var=StringVar()
    mobile_var=StringVar()
    qunt_var =StringVar() 
    sellmrp_var = StringVar()
    cust_name_entry= Entry(root,textvariable=name_var,font=Entry_Font1,width=20)
    cust_number_entry= Entry(root,textvariable=mobile_var,font=Entry_Font1,width=10)
    quntity_entry = Entry(root, textvariable = qunt_var,font=Entry_Font1,width=10)
    sellmrp_entry = Entry(root, textvariable = sellmrp_var,font=Entry_Font1,width=10)
    cust_name_entry.place(x = 950, y =350)
    cust_number_entry.place(x = 950, y =400)
    quntity_entry.place(x = 950, y =450)
    sellmrp_entry.place(x = 950, y =500)

    sell_button = Button(root, text = "Sell",style = Button_Style)
    newstock=0
    profit=0

    sheet9 = pds.read_excel(file9, sheet_name = "tube") 
    mainframe = Frame(root)
    mainframe.place(x=300,y=290)
    tkvar = StringVar(root)
    choices = sheet9[TYRE_SIZE].tolist()
    popupMenu = OptionMenu(mainframe, tkvar, *choices)
    popupMenu.configure(width=22)
    popupMenu.pack()

    x = datetime.datetime.now()
    date1=x.strftime("%d/%m/%y")
    date2=x.strftime("%d/%m/%Y")

    my_list1= StringVar(root)

    stock1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
    ndp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
    mrp1 = Label(text="######",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
    sell_fail = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
    sell_fail1 = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)
    sell_sucess = Label(text="############",font=Label_Font1,background=BG_WHITE,foreground=BG_WHITE)

    stock1.place(x = 450, y =350)
    ndp1.place(x = 450, y =430)
    mrp1.place(x = 450, y =510)

    sell_fail1.place(x=880,y=550)
    sell_fail.place(x=880,y=550)
    sell_sucess.place(x=870,y=550)

        
    def print1(var1,var2,var3):
        document = Document()
        ab=document.add_paragraph()
        ab.add_run('\t\t\t\t\tINVOICE').bold = True
        ab=document.add_paragraph()
        ab.add_run('Morya Sales & Services').bold = True
        
        a = document.add_paragraph('Main Road, Mane Nagar, Rendal-Hupari\n')
        a.add_run('Tal-Hatkanangle, Dist.-Kolhapur\n')
        a.add_run('Pin-416203\n')
        a.add_run('Phone: 8208600074, 9067173171\n')
        a.add_run('_________________________________________________________________________________________________________').bold = True

        name=str(cust_name_entry.get())
        mono=str(cust_number_entry.get())
        b = document.add_paragraph('\t\t\t\t\t\t\t\t\tDate : ')
        b.add_run(date2)
        b = document.add_paragraph('Customer Name : ')
        b.add_run(name)
        b = document.add_paragraph('Mobile No. : ')
        b.add_run(mono)

        print_total=var2*var3

        recordset = [
            {
                "No." : 1,
                "desc": tkvar.get(),
                "qty": var2,
                "Unit_Price": var3,
                "Total":print_total
            }
        ]
        table = document.add_table(rows=1, cols=5, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Quantity'
        hdr_cells[3].text = 'Unit Price'
        hdr_cells[4].text = 'Total'
        for item in recordset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['No.'])
            row_cells[1].text = str(item['desc'])
            row_cells[2].text = str(item['qty'])
            row_cells[3].text = str(item['Unit_Price'])
            row_cells[4].text = str(item['Total'])

        ac=document.add_paragraph()
        ac.add_run('\n\n\n\n\n\n\n\n\n\n\n\n\t\t\t\t\t\t\t\t   Morya Sales & Services').bold = True

        document.save('simple.docx')
        os.startfile('simple.docx','print')
        print('print')
 
    def sell(var1,var2):
        print('sell')
        sell_fail1 = Label(text = 'Enter Quntity', font=Label_Font4,background=BG_WHITE, foreground = FG_Red)
        sell_fail = Label(text = 'Enter Selling Price', font=Label_Font4,background=BG_WHITE, foreground = FG_Red)
        sell_sucess = Label(text = 'Successful', font=Label_Font5,background=BG_WHITE, foreground = FG_Green)
        if quntity_entry.get() == "":
            sell_fail1.place(x=880,y=550)
        elif sellmrp_entry.get() == "":
            sell_fail.place(x=880,y=550)
        else:
            sell_quntity = int(quntity_entry.get())
            sell_price = int(sellmrp_entry.get())
            name1=str(cust_name_entry.get())
            number1=int(cust_number_entry.get())
            newstock=var1-sell_quntity
            profit = (sell_price-var2)*sell_quntity
            print(sell_quntity)
            print(sell_price)
            print(newstock)
            print(profit)
            sell_fail.destroy()
            sell_fail1.destroy()
            sell_sucess.place(x=870,y=550)
        return newstock,profit,sell_quntity,sell_price

    def change_dropdown(*args):
        sheet9.set_index(TYRE_SIZE, inplace = True)
        my_list2=sheet9.loc[tkvar.get()][1]
        my_list3=sheet9.loc[tkvar.get()][2]
        my_list4=sheet9.loc[tkvar.get()][3]
        print (my_list2 , my_list3, my_list4)
        stock1 = Label(text=my_list4,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =350)
        ndp1 = Label(text=my_list2,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =430)
        mrp1 = Label(text=my_list3,font=Label_Font1,background=BG_WHITE,foreground=FG_Black).place(x = 450, y =510)
        def sell1():
            newstock,profit,sell_quntity,sell_price=sell(my_list4,my_list2)
            name1=str(cust_name_entry.get())
            number1=int(cust_number_entry.get())
            print('newstock=',newstock)
            sheet9.loc[tkvar.get(),'Stock']=newstock
            sheet9.to_excel(file9, sheet_name = 'tube')

            Sales_report_tube = pds.read_excel(file10, sheet_name = "Sheet1") 
            df = pds.DataFrame(Sales_report_tube)
            no=len(df['No.'])+1
            df2 = pds.DataFrame({'No.':[no],'Date':[date1],'Tube Size':[tkvar.get()],'Quantity':[sell_quntity],'NDP':[my_list2], 'MRP':[sell_price], 'Profit':[profit],"Customer Name":[name1],"Mobile No.":[number1]})
            append_sell =df.append(df2,ignore_index = True)
            append_sell.to_excel(file10, sheet_name = 'Sheet1', index=False) 
            def print2():
                print1(tkvar.get(),sell_quntity,sell_price)
            print_button = Button(root, text = "Print",style = Button_Style,command = print2)
            print_button.place(x = 880, y =640)
        sell_button = Button(root, text = 'Sell',style = Button_Style,command = sell1)
        sell_button.place(x = 880, y =600)
    tkvar.trace('w', change_dropdown)

def ebike():
    A = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
    A.place(x=0,y=195)
    B = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
    B.place(x=0,y=195)
    c = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1166,relief=FLAT) 
    c.place(x=202,y=195)
    header0 = Label(text="E-BIKE SECTION", font = Header_Font1,background=BG_DeepSkyBlue)
    header0.place(x=690,y=197)

def about():
    print('about')
    A = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
    A.place(x=0,y=195)
    # B = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
    # B.place(x=0,y=195)
    c = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1366,relief=FLAT) 
    # c.place(x=202,y=195)
    c.place(x=0,y=195)
    header0 = Label(text="ABOUT SECTION", font = Header_Font1,background=BG_DeepSkyBlue)
    header0.place(x=570,y=197)

    
    header1 = Label(text="MORYA MOTORS & MORYA SALES AND SERVICES",font=Label_Font1,background=BG_WHITE,foreground=FG_Black)
    header2=Label(text='Main Road, Mane Nagar, Rendal-Hupari,', font = Label_Font7,background=BG_WHITE)
    header3=Label(text='Tal-Hatkanangle, Dist.-Kolhapur', font = Label_Font7,background=BG_WHITE)
    header4=Label(text='Pin-416203', font = Label_Font7,background=BG_WHITE)
    header5=Label(text='Phone: 8208600074, 9067173171', font = Label_Font7,background=BG_WHITE)
    header1.place(x=50,y=270)
    header2.place(x=50,y=310)
    header3.place(x=50,y=340)
    header4.place(x=50,y=370)
    header5.place(x=50,y=400)

    d = Canvas(root, bg =BG_DeepSkyBlue, height = 15, width = 1366,relief=FLAT) 
    # d.place(x=202,y=450)
    d.place(x=0,y=450)
    header6 = Label(text="Software developed by,",font=("Helvetica", 15),background=BG_WHITE,foreground=FG_Black)
    header7=Label(text='Er.Onkar Anil Deshpande', font = Header_Font1,background=BG_WHITE)
    header9=Label(text='E-mail ID : onkard543@gmail.com', font = Label_Font7,background=BG_WHITE)
    header10=Label(text='Phone: 9552708854', font = Label_Font7,background=BG_WHITE)
    header6.place(x=880,y=490)
    header7.place(x=940,y=520)
    header9.place(x=940,y=560)
    header10.place(x=940,y=590)

def login():
    name=name_entry.get() 
    password=passw_entry.get() 
      
    print("The name is : " + name) 
    print("The password is : " + password) 

    if name == 'username':
        if password == '1234':
            login_succf.place(x = 1090,y = 350)
            print('welcome')
            main_frame1()           
        else:
            print('wrong entry')  
            login_fail.place(x = 1040,y = 350)
    else:
        print('wrong entry')  
        login_fail.place(x = 1040,y = 350)
    name_var.set("") 
    passw_var.set("")  
   
def main_frame():
    my.place(x=0,y=25)
    main_title3.place(x=980,y=80)
    main_title4.place(x=925,y=120)
    name_label.place(x = 950,y = 250)
    name_entry.place(x = 1080,y = 250) 
    passw_label.place(x = 950,y = 300)
    passw_entry.place(x = 1080,y = 300)
    login_button.place(x = 1050, y = 380)

def main_frame1():
    my.destroy()
    main_title3.destroy()
    main_title4.destroy()
    login_fail.destroy()
    passw_label.destroy()
    name_label.destroy()
    name_entry.destroy()
    passw_entry.destroy()
    login_button.destroy()
    login_succf.destroy()
    print('now i am in main frame')

    A = Canvas(root, bg =BG_LightSkyBlue, height = 195, width = 1366,relief=FLAT) 
    A.place(x=0,y=0)
    main_title1 = Label(text="MORYA MOTORS &",font=Label_Font8,background=BG_LightSkyBlue,foreground=FG_Red)
    main_title2 = Label(text="MORYA SALES AND SERVICES",font=Label_Font9,background=BG_LightSkyBlue,foreground=FG_Red)
    tyre_button = Button(root, text = "Tyres",style = Button_Style,command = tyre)
    ebike_button = Button(root, text = "Hero E-Bike",style = Button_Style,command = ebike)
    about_button = Button(root, text = "About",style = Button_Style,command = about)
    tubes_button = Button(root, text = "Tubes",style = Button_Style,command = tubes)
    update_button = Button(root, text = "Update",style = Button_Style,command = update1)
    main_title1.place(x=500,y=10)
    main_title2.place(x=450,y=60)
    tyre_button.place(x = 320, y =120)
    tubes_button.place(x = 520, y =120)
    # ebike_button.place(x = 720, y =120)
    update_button.place(x = 720, y =120)
    about_button.place(x = 920, y =120)
    tyre()

def update1():
    A = Canvas(root, bg =BG_WHITE, height = 500, width = 1366,relief=FLAT) 
    A.place(x=0,y=195)
    B = Canvas(root, bg =BG_DeepSkyBlue, height = 500, width = 200,relief=FLAT) 
    B.place(x=0,y=195)
    c = Canvas(root, bg =BG_DeepSkyBlue, height = 45, width = 1166,relief=FLAT) 
    c.place(x=202,y=195)
    header0 = Label(text="Data Update To GOOGLE SHEET ", font = Header_Font1,background=BG_DeepSkyBlue)
    header0.place(x=570,y=197)
    gc = gspread.service_account(filename='Credentials.json')
    Processing = Label(text="Processing..", font = Label_Font2,background=BG_WHITE,foreground = FG_Red)
    Successful = Label(text="Successful...", font = Label_Font2,background=BG_WHITE,foreground="Green")
    # white1= Label(text="###############", font = Label_Font2,background=BG_WHITE,foreground=BG_WHITE)
    # # Processing.place(x=540,y=300)
    # # Successful.place(x=540,y=300)
    # white1.place(x=450,y=400)
    # white1.place(x=940,y=400)

    def tyre_update():
        print('tyre update')
        A = Canvas(root, bg =BG_WHITE, height = 500, width = 1166,relief=FLAT) 
        A.place(x=201,y=241)
        def tyre_stock_update():
            print("tyre_stock_update")
            Processing.place(x=450,y=400)
            time.sleep(2)
                
            def CEAT_gs():
                file1=('Excel_files\\CEAT TYRE.xlsx')
                sheet1 = pds.read_excel(file1, sheet_name = CEAT_TYRE , dtype=str) 
                CEAT = gc.open_by_key("1h313-U8TQUETKi33AyLoqRhRy_SI2hBeATZBUIdJPCA")
                worksheet2 = CEAT.sheet1
                worksheet2.clear()
                colomnss=sheet1.columns.values.tolist()
                print(colomnss)
                worksheet2.append_row(colomnss)
                (row, col) = sheet1.shape
                for i in range(0, row):
                    user = sheet1.loc[i, :].tolist()
                    print(user)
                    worksheet2.append_row(user)

            def APOLLO_gs():
                file2=('Excel_files\\APOLLO TYRE.xlsx')
                sheet2 = pds.read_excel(file2, sheet_name = APOLLO_TYRE , dtype=str)
                APOLLO = gc.open_by_key("1gJMzDZiy_xQsuweMd9NaI4r2JVPa2oWRpf22cwzSjX4")
                worksheet3 = APOLLO.sheet1
                worksheet3.clear()
                colomnss=sheet2.columns.values.tolist()
                print(colomnss)
                worksheet3.append_row(colomnss)
                (row, col) = sheet2.shape
                for i in range(0, row):
                    user = sheet2.loc[i, :].tolist()
                    print(user)
                    worksheet3.append_row(user)

            def NUEMEX_gs():
                file3=('Excel_files\\NUEMEX TYRE.xlsx')
                sheet3 = pds.read_excel(file3, sheet_name = NUEMEX_TYRE , dtype=str)
                NUEMEX = gc.open_by_key("1CatZcYGQZHlGuf1Rdvosf1e22BpofU5IHIBRB6rrHRg")
                worksheet4 = NUEMEX.sheet1
                worksheet4.clear()
                colomnss=sheet3.columns.values.tolist()
                print(colomnss)
                worksheet4.append_row(colomnss)
                (row, col) = sheet3.shape
                for i in range(0, row):
                    user = sheet3.loc[i, :].tolist()
                    print(user)
                    worksheet4.append_row(user)

            def TVS_gs():
                file4=('Excel_files\\TVS TYRE.xlsx')
                sheet4 = pds.read_excel(file4, sheet_name = TVS_TYRE , dtype=str)
                TVS = gc.open_by_key("14sRAaLZwz1ubx1ak-f2WdP7qPqGAeBnlImWsS6oQ_6w")
                worksheet5 = TVS.sheet1
                worksheet5.clear()
                colomnss=sheet4.columns.values.tolist()
                print(colomnss)
                worksheet5.append_row(colomnss)
                (row, col) = sheet4.shape
                for i in range(0, row):
                    user = sheet4.loc[i, :].tolist()
                    print(user)
                    worksheet5.append_row(user)

            def STELLBIRD_gs():
                file5=('Excel_files\\STELLBIRID TYRE.xlsx')
                sheet5 = pds.read_excel(file5, sheet_name = STELLBIRID_TYRE , dtype=str)
                STELLBIRID = gc.open_by_key("1KrdNPUzoxSlVhPduOG6sGKH1EmVxlEkW0ZQrALQoCQg")
                worksheet6 = STELLBIRID.sheet1
                worksheet6.clear()
                colomnss=sheet5.columns.values.tolist()
                print(colomnss)
                worksheet6.append_row(colomnss)
                (row, col) = sheet5.shape
                for i in range(0, row):
                    user = sheet5.loc[i, :].tolist()
                    print(user)
                    worksheet6.append_row(user)

            def MRF_gs():
                file6=('Excel_files\\MRF TYRE.xlsx')
                sheet6 = pds.read_excel(file6, sheet_name = MRF_TYRE , dtype=str)
                MRF = gc.open_by_key("1HxJjn61M7afAYV4rwxxx4roCHED4pMtYIHM2IzqblDQ")
                worksheet7 = MRF.sheet1
                worksheet7.clear()
                colomnss=sheet6.columns.values.tolist()
                print(colomnss)
                worksheet7.append_row(colomnss)
                (row, col) = sheet6.shape
                for i in range(0, row):
                    user = sheet6.loc[i, :].tolist()
                    print(user)
                    worksheet7.append_row(user)

            def METRO_gs():
                file7=('Excel_files\\METRO TYRE.xlsx') 
                sheet7 = pds.read_excel(file7, sheet_name = METRO_TYRE , dtype=str)
                METRO = gc.open_by_key("14k_KDa_ksayEeMxcLaaNFy-Gfz2a0JxKreTW4cvEEBs")
                worksheet8 = METRO.sheet1
                worksheet8.clear()
                colomnss=sheet7.columns.values.tolist()
                print(colomnss)
                worksheet8.append_row(colomnss)
                (row, col) = sheet7.shape
                for i in range(0, row):
                    user = sheet7.loc[i, :].tolist()
                    print(user)
                    worksheet8.append_row(user)
            
            CEAT_gs()
            time.sleep(120)
            APOLLO_gs()
            time.sleep(120)
            NUEMEX_gs()
            TVS_gs()
            STELLBIRD_gs()
            time.sleep(120)
            MRF_gs()
            time.sleep(120)
            METRO_gs()
            time.sleep(120)

            Successful.place(x=450,y=400)

        def tyre_sale_update():
            print("tyre_sale_update")
            Processing.place(x=950,y=400)
            def Sales_report_excel_gs():
                file8=('Excel_files\\Sales_report.xlsx')
                Sales_report_excel = pds.read_excel(file8, sheet_name = "Sheet1" , dtype=str) 
                Sales_report = gc.open_by_key("1xvBayjKcxmRGMcwQVHqsJTK3kVR1IFIasvM8m_SfNLo")
                worksheet9 = Sales_report.sheet1
                worksheet9.clear()
                colomnss=Sales_report_excel.columns.values.tolist()
                print(colomnss)
                worksheet9.append_row(colomnss)
                (row, col) = Sales_report_excel.shape
                for i in range(0, row):
                    user = Sales_report_excel.loc[i, :].tolist()
                    print(user)
                    worksheet9.append_row(user)
            
            Sales_report_excel_gs()
            time.sleep(120)
            Successful.place(x=950,y=400)

        tyre_stock_title1 = Label(text=  "TYRE STOCK UPDATE", font = Label_Font2,background=BG_WHITE).place(x=400,y=350)
        tyre_sale_title1 = Label(text=  "TYRE SALE UPDATE", font = Label_Font2,background=BG_WHITE).place(x=900,y=350)
        tyre_stock_update_button = Button(text = "UPDATE",style = Button_Style,command = tyre_stock_update).place(x=450,y=450)
        tyre_sale_update_button = Button(text = "UPDATE",style = Button_Style,command = tyre_sale_update).place(x=950,y=450)
        
    def tube_update():
        print('tube update')
        A = Canvas(root, bg =BG_WHITE, height = 500, width = 1166,relief=FLAT) 
        A.place(x=201,y=241)
        def tube_stock_update():
            print("tube_stock_update")
            Processing.place(x=450,y=400)
            time.sleep(2)         
            def tube_gs():
                file10=('Excel_files\\tube.xlsx') 
                tube_excel = pds.read_excel(file10, sheet_name = "tube", dtype=str)
                tube = gc.open_by_key("1TkgGF9hOqPrcQ0bySSQI4YU1cZmzwiq4d2cOco7YtJg")
                worksheet1 = tube.sheet1
                worksheet1.clear()
                colomnss=tube_excel.columns.values.tolist()
                print(colomnss)
                worksheet1.append_row(colomnss)
                (row, col) = tube_excel.shape
                for i in range(0, row):
                    user = tube_excel.loc[i, :].tolist()
                    print(user)
                    worksheet1.append_row(user)
            tube_gs()
            time.sleep(120)
            Successful.place(x=450,y=400)

        def tube_sale_update():
            print("tube_sale_update")
            Processing.place(x=950,y=400)
            def Sales_report_tube_excel_gs():
                file9=('Excel_files\\Sales_report_tube.xlsx')
                Sales_report_tube_excel = pds.read_excel(file9, sheet_name = "Sheet1" , dtype=str) 
                # df2 = pds.DataFrame(Sales_report_tube_excel)
                Sales_report_tube = gc.open_by_key("1O-prfLxTcT7AzEktNMln5QFFdeFYVFFSCcbUk6G6pZc")
                worksheet10 = Sales_report_tube.sheet1
                worksheet10.clear()
                colomnss=Sales_report_tube_excel.columns.values.tolist()
                print(colomnss)
                worksheet10.append_row(colomnss)
                (row, col) = Sales_report_tube_excel.shape
                for i in range(0, row):
                    user = Sales_report_tube_excel.loc[i, :].tolist()
                    print(user)
                    worksheet10.append_row(user)

            Sales_report_tube_excel_gs()
            time.sleep(120)
            Successful.place(x=950,y=400)

        tube_stock_title1 = Label(text=  "TUBE STOCK UPDATE", font = Label_Font2,background=BG_WHITE).place(x=400,y=350)
        tube_sale_title1 = Label(text=  "TUBE SALE UPDATE", font = Label_Font2,background=BG_WHITE).place(x=900,y=350)
        tube_stock_update_button = Button(text = "UPDATE",style = Button_Style,command = tube_stock_update).place(x=450,y=450)
        tube_sale_update_button = Button(text = "UPDATE",style = Button_Style,command = tube_sale_update).place(x=950,y=450)
        
    tyre_update_btn = Button(root, text = "TYRE",style = Button_Style,command=tyre_update)
    tyre_update_btn.place(x=40,y=245)
    tube_update_btn = Button(root, text = "TUBE",style = Button_Style,command=tube_update)
    tube_update_btn.place(x=40,y=300)
    tyre_update()
 
try:
    #create a logger
    logger = logging.getLogger(Root_Title)
    #set logging level
    logger.setLevel(logging.DEBUG)

    handler = logging.FileHandler(Debug_log)
    # create a logging format
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

    root=Tk()
    root.geometry(Root_Geometry)
    root.title(Root_Title) 
    root.configure(background=BG_WHITE)

    name_var =StringVar() 
    passw_var = StringVar()

    style1 = Style()
    style1.configure(Button_Style, font = Label_Font3,background=BG_LightSkyBlue, foreground = FG_Black)

    main_title3 = Label(text="MORYA MOTORS & ",font=Label_Font9,background=BG_WHITE,foreground=FG_Red)
    main_title4 = Label(text="MORYA SALES AND SERVICES",font=Label_Font1,background=BG_WHITE,foreground=FG_Red)

    name_label = Label(root, text = 'Username', font=Label_Font3,background=BG_WHITE, foreground = FG_Black) 
    name_entry = Entry(root, textvariable = name_var,font=Entry_Font1)
    passw_label = Label(root,text = 'Password', font = Label_Font3,background=BG_WHITE, foreground = FG_Black)
    passw_entry= Entry(root, textvariable = passw_var, font = Entry_Font1, show = '*') 
    login_succf = Label(root, text = 'Login Succesfully', font=Label_Font4,background=BG_WHITE, foreground =FG_Green) 
    login_fail = Label(root, text = 'Invalid Username or Password', font=Label_Font4,background=BG_WHITE, foreground =FG_Red)
    login_button = Button(root, text = "Login",style = Button_Style,command = login)

    simg = ImageTk.PhotoImage(Image.open(ImagePath))
    my = Label(root,image=simg)
    my.image = simg

    main_frame()

    root.mainloop()

except Exception as e:
    logger.debug(STATUS_ERROR1 + str(e))

